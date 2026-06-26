"""
Legora AI — Document Parser.

Unified interface for extracting text from PDF, DOCX, TXT, and scanned files.
"""

import io
import logging
from dataclasses import dataclass, field

import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    """A single page of extracted text."""

    page_number: int
    text: str
    is_ocr: bool = False


@dataclass
class ParsedDocument:
    """Result of document parsing."""

    pages: list[ParsedPage] = field(default_factory=list)
    full_text: str = ""
    page_count: int = 0
    detected_language: str | None = None
    ocr_applied: bool = False
    extraction_quality: float = 1.0

    @property
    def total_chars(self) -> int:
        return len(self.full_text)


def parse_pdf(file_content: bytes) -> ParsedDocument:
    """
    Parse a PDF file. Extract text per page.
    Falls back to OCR if a page has very little text.
    """
    pages: list[ParsedPage] = []
    ocr_applied = False

    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text").strip()

            # If page has very little text, it might be scanned — try OCR
            if len(text) < 50:
                try:
                    ocr_text = _ocr_page(page)
                    if ocr_text and len(ocr_text) > len(text):
                        text = ocr_text
                        ocr_applied = True
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num + 1}: {e}")

            pages.append(ParsedPage(
                page_number=page_num + 1,
                text=text,
                is_ocr=ocr_applied,
            ))
        doc.close()
    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        raise ValueError(f"Failed to parse PDF: {e}")

    full_text = "\n\n".join(p.text for p in pages if p.text)
    quality = _estimate_quality(full_text, len(pages))

    return ParsedDocument(
        pages=pages,
        full_text=full_text,
        page_count=len(pages),
        ocr_applied=ocr_applied,
        extraction_quality=quality,
    )


def parse_docx(file_content: bytes) -> ParsedDocument:
    """Parse a DOCX file. Extract text from paragraphs and tables."""
    try:
        doc = DocxDocument(io.BytesIO(file_content))
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)

    # DOCX doesn't have native page numbers, treat as single page
    pages = [ParsedPage(page_number=1, text=full_text)]

    return ParsedDocument(
        pages=pages,
        full_text=full_text,
        page_count=1,
        extraction_quality=_estimate_quality(full_text, 1),
    )


def parse_txt(file_content: bytes) -> ParsedDocument:
    """Parse a plain text file."""
    try:
        text = file_content.decode("utf-8")
    except UnicodeDecodeError:
        text = file_content.decode("latin-1")

    pages = [ParsedPage(page_number=1, text=text.strip())]

    return ParsedDocument(
        pages=pages,
        full_text=text.strip(),
        page_count=1,
        extraction_quality=1.0,
    )


def parse_document(file_content: bytes, mime_type: str) -> ParsedDocument:
    """
    Unified document parser. Routes to the appropriate parser based on MIME type.
    """
    parser_map = {
        "application/pdf": parse_pdf,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": parse_docx,
        "text/plain": parse_txt,
    }

    parser = parser_map.get(mime_type)
    if not parser:
        # Try to infer from content
        if file_content[:4] == b"%PDF":
            parser = parse_pdf
        elif file_content[:2] == b"PK":
            parser = parse_docx
        else:
            parser = parse_txt

    return parser(file_content)


def _ocr_page(page) -> str:
    """
    OCR a single PDF page using pytesseract.
    Renders the page as an image, then performs OCR.
    """
    try:
        import pytesseract
        from PIL import Image

        # Render page to image at 300 DPI
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = pytesseract.image_to_string(img)
        return text.strip()
    except ImportError:
        logger.warning("pytesseract not available, skipping OCR")
        return ""


def _estimate_quality(text: str, page_count: int) -> float:
    """
    Estimate text extraction quality as a score between 0 and 1.
    Based on character density, word count, and garbage character ratio.
    """
    if not text or page_count == 0:
        return 0.0

    word_count = len(text.split())
    char_count = len(text)

    # Average words per page (typical contract has 300-600 words per page)
    avg_words_per_page = word_count / page_count
    if avg_words_per_page < 20:
        return 0.2  # Very low — probably failed extraction

    # Check for garbage characters
    garbage_chars = sum(1 for c in text if ord(c) > 65535 or (ord(c) < 32 and c not in "\n\r\t"))
    garbage_ratio = garbage_chars / char_count if char_count > 0 else 0

    quality = min(1.0, avg_words_per_page / 300)
    quality *= (1 - garbage_ratio)
    return round(max(0.0, min(1.0, quality)), 2)
